#!/usr/bin/env python3
"""
document_parser.py — Extrator de texto de documentos PDF e DOCX.

Responsável por ler o conteúdo textual de documentos jurídicos
nos formatos PDF e DOCX, retornando texto limpo e estruturado.

Dependências:
    - pdfplumber (para PDFs)
    - python-docx (para DOCX)

Uso:
    python3 document_parser.py --input <arquivo> --output <saida.txt>
"""

import argparse
import os
import re
import sys
from pathlib import Path


class DocumentParser:
    """
    Parser de documentos jurídicos.
    Suporta PDF e DOCX, extraindo texto limpo com metadados básicos.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

    def parse(self, file_path: str) -> str:
        """
        Extrai texto de um documento PDF ou DOCX.

        Args:
            file_path: Caminho completo do arquivo

        Returns:
            Texto extraído do documento

        Raises:
            ValueError: Se o formato do arquivo não for suportado
            FileNotFoundError: Se o arquivo não existir
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")

        extension = path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Formato não suportado: {extension}. "
                f"Use: {self.SUPPORTED_EXTENSIONS}"
            )

        if extension == ".pdf":
            return self._parse_pdf(path)
        elif extension == ".docx":
            return self._parse_docx(path)

    def _parse_pdf(self, path: Path) -> str:
        """
        Extrai texto de um arquivo PDF usando pdfplumber.

        Args:
            path: Path do arquivo PDF

        Returns:
            Texto completo extraído de todas as páginas
        """
        try:
            import pdfplumber
        except ImportError:
            raise ImportError(
                "pdfplumber é necessário para ler PDFs. "
                "Instale com: pip install pdfplumber"
            )

        pages_text = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    # Adiciona marcador de página para referência
                    pages_text.append(f"[PÁGINA {i + 1}]\n{text}")

        raw_text = "\n\n".join(pages_text)
        return self._clean_text(raw_text)

    def _parse_docx(self, path: Path) -> str:
        """
        Extrai texto de um arquivo DOCX usando python-docx.

        Args:
            path: Path do arquivo DOCX

        Returns:
            Texto completo extraído de todos os parágrafos
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx é necessário para ler DOCX. "
                "Instale com: pip install python-docx"
            )

        doc = Document(str(path))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Identifica estilos de título para preservar estrutura
                style_name = para.style.name.lower() if para.style else ""
                if "heading" in style_name or "título" in style_name:
                    paragraphs.append(f"\n## {text}\n")
                else:
                    paragraphs.append(text)

        # Extrai também texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(f"[TABELA] {row_text}")

        raw_text = "\n".join(paragraphs)
        return self._clean_text(raw_text)

    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Limpa e normaliza o texto extraído.

        - Remove espaços múltiplos
        - Remove linhas em branco excessivas
        - Normaliza quebras de linha
        - Remove caracteres de controle

        Args:
            text: Texto bruto extraído

        Returns:
            Texto limpo e normalizado
        """
        # Remove caracteres de controle (exceto newline e tab)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

        # Normaliza espaços múltiplos em uma linha
        text = re.sub(r'[ \t]+', ' ', text)

        # Remove espaços no início e fim de cada linha
        lines = [line.strip() for line in text.split('\n')]

        # Remove linhas em branco excessivas (máximo 2 consecutivas)
        cleaned_lines = []
        blank_count = 0
        for line in lines:
            if not line:
                blank_count += 1
                if blank_count <= 2:
                    cleaned_lines.append(line)
            else:
                blank_count = 0
                cleaned_lines.append(line)

        return '\n'.join(cleaned_lines).strip()

    def get_metadata(self, file_path: str) -> dict:
        """
        Extrai metadados básicos do documento.

        Args:
            file_path: Caminho do arquivo

        Returns:
            Dicionário com metadados (nome, tamanho, formato, páginas)
        """
        path = Path(file_path)
        metadata = {
            "filename": path.name,
            "extension": path.suffix.lower(),
            "size_bytes": path.stat().st_size,
            "size_kb": round(path.stat().st_size / 1024, 2),
        }

        if path.suffix.lower() == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(str(path)) as pdf:
                    metadata["pages"] = len(pdf.pages)
            except Exception:
                metadata["pages"] = None

        return metadata


def main():
    """Ponto de entrada via linha de comando."""
    parser = argparse.ArgumentParser(
        description="Extrator de texto de documentos jurídicos"
    )
    parser.add_argument(
        "--input", "-i", required=True,
        help="Caminho do arquivo PDF ou DOCX"
    )
    parser.add_argument(
        "--output", "-o", default=None,
        help="Caminho do arquivo de saída (opcional)"
    )

    args = parser.parse_args()

    doc_parser = DocumentParser()

    # Extrai metadados
    metadata = doc_parser.get_metadata(args.input)
    print(f"Arquivo: {metadata['filename']}")
    print(f"Formato: {metadata['extension']}")
    print(f"Tamanho: {metadata['size_kb']} KB")
    if metadata.get("pages"):
        print(f"Páginas: {metadata['pages']}")

    # Extrai texto
    text = doc_parser.parse(args.input)
    print(f"Caracteres extraídos: {len(text)}")

    # Salva saída
    if args.output:
        Path(args.output).write_text(text, encoding="utf-8")
        print(f"Texto salvo em: {args.output}")
    else:
        print("\n--- TEXTO EXTRAÍDO ---\n")
        print(text[:2000])
        if len(text) > 2000:
            print(f"\n... [{len(text) - 2000} caracteres restantes]")


if __name__ == "__main__":
    main()
